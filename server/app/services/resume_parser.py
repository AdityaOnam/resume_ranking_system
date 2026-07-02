"""
Resume Parser v2.0
Combines best-of-breed techniques from both the existing resume-ranking-system parser
and the eightfold_transformer PDF extractor.

Key design decisions:
  - PyMuPDF (fitz) for PDF extraction — better layout and hyperlink support than PyPDF2.
  - python-docx for DOCX extraction.
  - Section-aware parsing — splits text into named sections, then runs targeted extraction
    within each section for higher precision.
  - Skill extraction uses the external skill_taxonomy.json for maintainability.
  - Skill canonicalization maps abbreviations ("js" → "JavaScript", "ml" → "Machine Learning").
  - Returns structured JSON matching the v2.0 database schema (parsed_data JSONB column).
"""

import os
import re
import json
import logging
from typing import Optional, List, Dict, Any

import fitz  # PyMuPDF
import spacy

from app.services.github_service import GitHubService
from app.services.llm_service import LLMService

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Load spaCy model (singleton, loaded once at module import)
# ---------------------------------------------------------------------------
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    logger.warning("Downloading spaCy model en_core_web_sm …")
    import subprocess, sys
    subprocess.run([sys.executable, "-m", "spacy", "download", "en_core_web_sm"], check=True)
    nlp = spacy.load("en_core_web_sm")

# ---------------------------------------------------------------------------
# Load skill taxonomy from JSON data file
# ---------------------------------------------------------------------------
_DATA_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "..", "data")

def _load_skill_taxonomy() -> Dict[str, List[str]]:
    path = os.path.join(_DATA_DIR, "skill_taxonomy.json")
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except FileNotFoundError:
        logger.warning(f"skill_taxonomy.json not found at {path}, using empty taxonomy.")
        return {}

SKILL_TAXONOMY: Dict[str, List[str]] = _load_skill_taxonomy()

# Build a flat lookup: lowercase skill name → canonical display name
_SKILL_LOOKUP: Dict[str, str] = {}
for _category, _skills in SKILL_TAXONOMY.items():
    for _skill in _skills:
        _SKILL_LOOKUP[_skill.lower()] = _skill

# Abbreviation / alias map → canonical name
SKILL_ALIASES: Dict[str, str] = {
    "js": "JavaScript",
    "ts": "TypeScript",
    "py": "Python",
    "python3": "Python",
    "cpp": "C++",
    "c/c++": "C++",
    "golang": "Go",
    "reactjs": "React",
    "react.js": "React",
    "vuejs": "Vue.js",
    "vue": "Vue.js",
    "angular.js": "Angular",
    "angularjs": "Angular",
    "node": "Node.js",
    "node.js": "Node.js",
    "nodejs": "Node.js",
    "express": "Express.js",
    "expressjs": "Express.js",
    "tf": "TensorFlow",
    "sklearn": "scikit-learn",
    "scikit learn": "scikit-learn",
    "opencv": "OpenCV",
    "postgres": "PostgreSQL",
    "mongo": "MongoDB",
    "k8s": "Kubernetes",
    "gcp": "Google Cloud Platform",
    "aws": "Amazon Web Services",
    "ml": "Machine Learning",
    "dl": "Deep Learning",
    "nlp": "Natural Language Processing",
    "cv": "Computer Vision",
    "ai": "Artificial Intelligence",
    "llm": "Large Language Models",
    "rl": "Reinforcement Learning",
    "oop": "Object-Oriented Programming",
    "dsa": "Data Structures",
    "dbms": "Database Management Systems",
    "os": "Operating Systems",
    "cn": "Computer Networks",
    "ci/cd": "CI/CD",
    "devops": "DevOps",
    "iot": "Internet of Things",
    "rpa": "Robotics",
    "seo": "SEO",
    "sem": "SEM",
}

# ---------------------------------------------------------------------------
# Section header patterns (from eightfold's approach, extended)
# ---------------------------------------------------------------------------
SECTION_PATTERNS: Dict[str, str] = {
    "contact":    r"(?i)^(?:contact\s*(?:info(?:rmation)?|details)?|personal\s*(?:info(?:rmation)?|details))$",
    "education":  r"(?i)^(?:education|academic\s*(?:profile|qualifications?|background)|qualifications?)$",
    "skills":     r"(?i)^(?:(?:technical\s+)?skills|technologies|tech\s*stack|core\s*competenc(?:ies|y)|programming\s*languages|tools\s*(?:and|&)\s*technologies)$",
    "experience": r"(?i)^(?:(?:work\s+|professional\s+)?experience|employment(?:\s*history)?|internship(?:s)?|work\s*history)$",
    "projects":   r"(?i)^(?:(?:academic\s+|technical\s+|personal\s+|selected\s+)?projects?|project\s*(?:work|experience)|portfolio)$",
    "achievements": r"(?i)^(?:achievements?|awards?|honors?|certifications?|publications?)$",
    "summary":    r"(?i)^(?:(?:career\s+|professional\s+)?summary|(?:career\s+)?objective|profile|about\s*me)$",
}

# ---------------------------------------------------------------------------
# Contact regex patterns
# ---------------------------------------------------------------------------
EMAIL_PATTERN = re.compile(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+")

PHONE_PATTERNS = [
    re.compile(r"\+?\d{1,3}[-.\s]?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}"),
    re.compile(r"\+?\d{10,13}"),
]

LINKEDIN_PATTERN = re.compile(r"(?:https?://)?(?:www\.)?linkedin\.com/in/[\w-]+", re.IGNORECASE)
GITHUB_PATTERN = re.compile(r"(?:https?://)?(?:www\.)?github\.com/[\w-]+", re.IGNORECASE)

# ---------------------------------------------------------------------------
# GPA / CPI patterns (from existing parser — battle-tested)
# ---------------------------------------------------------------------------
GPA_PATTERNS = [
    re.compile(r"(?:cgpa|cpi|gpa)\s*[:/\-]?\s*(\d+\.?\d*)\s*/\s*(\d+\.?\d*)", re.IGNORECASE),
    re.compile(r"(?:cgpa|cpi|gpa)\s*[:/\-]?\s*(\d+\.\d+)", re.IGNORECASE),
    re.compile(r"(\d+\.\d+)\s*/\s*10(?:\.\d+)?", re.IGNORECASE),
    re.compile(r"(\d+\.\d+)\s+out\s+of\s+(?:10|4)", re.IGNORECASE),
]

# ---------------------------------------------------------------------------
# Degree patterns (from eightfold, extended)
# ---------------------------------------------------------------------------
DEGREE_PATTERNS = [
    (re.compile(r"(?:B\.?\s*Tech\.?|Bachelor\s*of\s*Technology)", re.IGNORECASE), "B.Tech"),
    (re.compile(r"(?:M\.?\s*Tech\.?|Master\s*of\s*Technology)", re.IGNORECASE), "M.Tech"),
    (re.compile(r"(?:B\.?\s*E\.?|Bachelor\s*of\s*Engineering)", re.IGNORECASE), "B.E."),
    (re.compile(r"(?:M\.?\s*E\.?|Master\s*of\s*Engineering)", re.IGNORECASE), "M.E."),
    (re.compile(r"(?:B\.?\s*Sc\.?|Bachelor\s*of\s*Science)", re.IGNORECASE), "B.Sc"),
    (re.compile(r"(?:M\.?\s*Sc\.?|M\.?\s*S\.?|Master\s*of\s*Science)", re.IGNORECASE), "M.Sc"),
    (re.compile(r"(?:Ph\.?\s*D\.?|Doctorate)", re.IGNORECASE), "Ph.D"),
    (re.compile(r"(?:MBA|Master\s*of\s*Business)", re.IGNORECASE), "MBA"),
    (re.compile(r"(?:B\.?\s*A\.?|Bachelor\s*of\s*Arts)", re.IGNORECASE), "B.A."),
    (re.compile(r"(?:B\.?\s*Com\.?|Bachelor\s*of\s*Commerce)", re.IGNORECASE), "B.Com"),
    (re.compile(r"(?:BCA|Bachelor\s*of\s*Computer\s*Applications)", re.IGNORECASE), "BCA"),
    (re.compile(r"(?:MCA|Master\s*of\s*Computer\s*Applications)", re.IGNORECASE), "MCA"),
]

# Field-of-study patterns
FIELD_PATTERNS = re.compile(
    r"(?:in|of)\s+"
    r"(Mathematics\s*(?:and|&)\s*Computing"
    r"|Computer\s*Science(?:\s*(?:and|&)\s*Engineering)?"
    r"|Electrical\s*Engineering"
    r"|Electronics\s*(?:and|&)\s*Communication(?:\s*Engineering)?"
    r"|Mechanical\s*Engineering"
    r"|Civil\s*Engineering"
    r"|Chemical\s*Engineering"
    r"|Petroleum\s*Engineering"
    r"|Aerospace\s*Engineering"
    r"|Information\s*Technology"
    r"|Data\s*Science"
    r"|Artificial\s*Intelligence"
    r"|Biomedical\s*Engineering"
    r"|Environmental\s*Engineering"
    r"|Automobile\s*Engineering"
    r"|Metallurgical\s*Engineering"
    r")",
    re.IGNORECASE
)

# Institution patterns
INSTITUTION_PATTERNS = [
    re.compile(r"((?:Indian Institute of Technology|IIT|NIT|IIIT|BITS|VIT|SRM|DTU|NSUT|IISC|ISI)[^\n,;]*)", re.IGNORECASE),
    re.compile(r"((?:University|Institute|College|School|Academy)[^\n,;]{3,60})", re.IGNORECASE),
]

# Year extraction
YEAR_RANGE_PATTERN = re.compile(r"(20\d{2})\s*[-–—]\s*(20\d{2}|Present|Current|Ongoing)", re.IGNORECASE)
SINGLE_YEAR_PATTERN = re.compile(r"(?:expected|graduating|graduation|batch)\s*(?:in|of|:|-|–)?\s*(20\d{2})", re.IGNORECASE)


# ===========================================================================
# Main Parser Class
# ===========================================================================
class ResumeParser:
    """
    Parses PDF and DOCX resumes into structured JSON.

    Output schema (matches Supabase resumes.parsed_data):
    {
        "contact": {"name", "email", "phone", "linkedin", "github"},
        "education": [{"institution", "degree", "field", "gpa", "start_year", "end_year"}],
        "skills": ["Python", "React", ...],
        "experience": [{"company", "role", "duration", "description"}],
        "projects": [{"title", "description", "technologies"}]
    }
    """

    # ----- Public API -----

    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        Main entry point.  Accepts a PDF or DOCX file path.
        Returns structured parsed_data dict + raw_text.
        """
        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".pdf":
            raw_text, embedded_links = self._extract_text_from_pdf(file_path)
        elif ext in (".docx", ".doc"):
            raw_text = self._extract_text_from_docx(file_path)
            embedded_links = []
        else:
            raise ValueError(f"Unsupported file type: {ext}. Only PDF and DOCX are supported.")

        if not raw_text or len(raw_text.strip()) < 50:
            return {
                "raw_text": raw_text or "",
                "parsed_data": {"error": "Could not extract meaningful text from file."},
            }

        # Split into named sections
        sections = self._split_into_sections(raw_text)

        contact = self._extract_contact(raw_text, sections, embedded_links)
        education = self._extract_education(raw_text, sections)
        skills = self._extract_skills(raw_text, sections)
        experience = self._extract_experience(raw_text, sections)
        projects = self._extract_projects(raw_text, sections)

        # 2. GitHub API Extraction
        github_stats = None
        github_url = contact.get("github")
        if github_url:
            logger.info(f"GitHub URL found: {github_url}. Fetching data...")
            github_svc = GitHubService()
            github_stats = github_svc.extract_github_data(github_url, projects)
            # Merge new skills found in READMEs
            if not github_stats.get("error"):
                readme_skills = github_stats.get("readme_skills", [])
                skills = list(set(skills + readme_skills))

        # 3. LLM Structured Extraction (Fallback & Enrichment)
        logger.info("Triggering Llama 3.1 LLM extraction...")
        llm_svc = LLMService()
        llm_data = llm_svc.extract_resume_data(raw_text)

        # 4. Merge results
        if llm_data:
            # LLM is better at structuring text but regex is safer for exact contacts
            llm_skills = llm_data.get("skills", [])
            merged_skills = list(set(skills + [s.title() for s in llm_skills if isinstance(s, str)]))

            parsed_data = {
                "contact": contact,  # Prefer heuristic for exact matching
                "education": llm_data.get("education") or education,
                "skills": sorted(merged_skills),
                "experience": llm_data.get("experience") or experience,
                "projects": llm_data.get("projects") or projects,
            }
        else:
            parsed_data = {
                "contact": contact,
                "education": education,
                "skills": sorted(skills),
                "experience": experience,
                "projects": projects,
            }

        # Attach github stats if available
        if github_stats and not github_stats.get("error"):
            parsed_data["github_stats"] = github_stats

        return {
            "raw_text": raw_text,
            "parsed_data": parsed_data,
        }

    # ------------------------------------------------------------------
    # Text Extraction
    # ------------------------------------------------------------------

    def _extract_text_from_pdf(self, pdf_path: str) -> tuple:
        """Extract text and embedded hyperlinks from PDF using PyMuPDF."""
        try:
            doc = fitz.open(pdf_path)
            pages_text = []
            embedded_links = []

            for page in doc:
                pages_text.append(page.get_text("text"))

                # Extract clickable hyperlinks (catches GitHub/LinkedIn icons)
                for link in page.get_links():
                    uri = link.get("uri")
                    if uri:
                        embedded_links.append(uri)

            doc.close()
            return "\n".join(pages_text), embedded_links
        except Exception as e:
            logger.error(f"Error extracting text from PDF {pdf_path}: {e}")
            return "", []

    def _extract_text_from_docx(self, docx_path: str) -> str:
        """Extract text from DOCX using python-docx."""
        try:
            from docx import Document
            doc = Document(docx_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract text from tables (many resumes use tables for layout)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            paragraphs.append(text)
            return "\n".join(paragraphs)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {docx_path}: {e}")
            return ""

    # ------------------------------------------------------------------
    # Section Splitting (from eightfold approach)
    # ------------------------------------------------------------------

    def _split_into_sections(self, text: str) -> Dict[str, str]:
        """
        Splits resume text into named sections by matching common header patterns.
        Returns dict like {"skills": "...", "experience": "...", ...}
        """
        lines = text.split("\n")
        boundaries: List[tuple] = []  # (line_index, section_name)

        for i, line in enumerate(lines):
            stripped = line.strip()
            if not stripped or len(stripped) > 80:
                continue
            for section_name, pattern in SECTION_PATTERNS.items():
                if re.match(pattern, stripped):
                    boundaries.append((i, section_name))
                    break

        if not boundaries:
            return {"unknown": text}

        boundaries.sort(key=lambda x: x[0])

        sections: Dict[str, str] = {}
        for idx, (start_line, section_name) in enumerate(boundaries):
            end_line = boundaries[idx + 1][0] if idx + 1 < len(boundaries) else len(lines)
            section_text = "\n".join(lines[start_line + 1 : end_line])
            if section_name in sections:
                sections[section_name] += "\n" + section_text
            else:
                sections[section_name] = section_text

        return sections

    # ------------------------------------------------------------------
    # Contact Extraction
    # ------------------------------------------------------------------

    def _extract_contact(self, text: str, sections: Dict[str, str], embedded_links: List[str]) -> Dict[str, Any]:
        """Extract name, email, phone, LinkedIn, GitHub."""
        # Name: first non-empty line is almost always the name
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        name = lines[0] if lines else ""
        # Sanitize: if "name" looks like an email or URL, skip it
        if "@" in name or "http" in name.lower() or len(name) > 60:
            name = ""

        # Email
        email_match = EMAIL_PATTERN.search(text)
        email = email_match.group(0) if email_match else None

        # Phone
        phone = None
        for pat in PHONE_PATTERNS:
            m = pat.search(text)
            if m:
                phone = re.sub(r"[^\d+]", "", m.group(0))
                if len(phone) >= 10:
                    break
                phone = None

        # LinkedIn & GitHub — check both text and embedded PDF links
        all_text = text + " " + " ".join(embedded_links)
        linkedin_match = LINKEDIN_PATTERN.search(all_text)
        github_match = GITHUB_PATTERN.search(all_text)

        linkedin = linkedin_match.group(0) if linkedin_match else None
        github = github_match.group(0) if github_match else None
        if linkedin and not linkedin.startswith("http"):
            linkedin = "https://" + linkedin
        if github and not github.startswith("http"):
            github = "https://" + github

        return {
            "name": name,
            "email": email,
            "phone": phone,
            "linkedin": linkedin,
            "github": github,
        }

    # ------------------------------------------------------------------
    # Education Extraction
    # ------------------------------------------------------------------

    def _extract_education(self, text: str, sections: Dict[str, str]) -> List[Dict[str, Any]]:
        """Extract education entries with institution, degree, field, GPA, years."""
        edu_text = sections.get("education", "")
        if not edu_text:
            edu_text = text  # fallback: search entire document

        entries = []

        # --- GPA extraction ---
        gpa = None
        for pat in GPA_PATTERNS:
            m = pat.search(edu_text)
            if m:
                try:
                    gpa = float(m.group(1))
                    # Normalize if it's on a 4.0 scale
                    if gpa <= 4.0 and "/4" in edu_text:
                        gpa = round(gpa * 2.5, 2)
                    break
                except (ValueError, IndexError):
                    continue

        # --- Institution ---
        institution = None
        for pat in INSTITUTION_PATTERNS:
            m = pat.search(edu_text)
            if m:
                institution = m.group(1).strip()
                break

        # --- Degree ---
        degree = None
        for pat, deg_name in DEGREE_PATTERNS:
            if pat.search(edu_text):
                degree = deg_name
                break

        # --- Field of study ---
        field = None
        field_match = FIELD_PATTERNS.search(edu_text)
        if field_match:
            field = field_match.group(1).strip()

        # --- Years ---
        start_year, end_year = None, None
        yr = YEAR_RANGE_PATTERN.search(edu_text)
        if yr:
            start_year = int(yr.group(1))
            end_str = yr.group(2)
            end_year = None if end_str.lower() in ("present", "current", "ongoing") else int(end_str)
        else:
            yr2 = SINGLE_YEAR_PATTERN.search(edu_text)
            if yr2:
                end_year = int(yr2.group(1))

        if institution or degree:
            entry: Dict[str, Any] = {
                "institution": institution or "Unknown",
                "degree": degree,
                "field": field,
                "start_year": start_year,
                "end_year": end_year,
            }
            if gpa is not None:
                entry["gpa"] = gpa
            entries.append(entry)

        return entries

    # ------------------------------------------------------------------
    # Skills Extraction
    # ------------------------------------------------------------------

    def _extract_skills(self, text: str, sections: Dict[str, str]) -> List[str]:
        """
        Extract skills by matching against the skill taxonomy.
        Searches both the Skills section and the full text.
        Returns a deduplicated, canonicalized list.
        """
        found: Dict[str, str] = {}  # lowercase → canonical name

        # Prioritize the skills section, but also scan full text
        search_texts = []
        if "skills" in sections:
            search_texts.append(sections["skills"])
        search_texts.append(text)

        combined = " ".join(search_texts).lower()

        # 1. Match against taxonomy (exact word boundary match)
        for lower_name, canonical in _SKILL_LOOKUP.items():
            # For very short names (<=2 chars like "C", "R", "Go"), require stricter context
            if len(lower_name) <= 2:
                # Only match if preceded/followed by non-alpha or in a comma-separated list
                pattern = r"(?<![a-zA-Z])" + re.escape(lower_name) + r"(?![a-zA-Z])"
            else:
                pattern = r"\b" + re.escape(lower_name) + r"\b"

            if re.search(pattern, combined):
                found[lower_name] = canonical

        # 2. Match aliases
        for alias, canonical in SKILL_ALIASES.items():
            if len(alias) <= 2:
                pattern = r"(?<![a-zA-Z])" + re.escape(alias) + r"(?![a-zA-Z])"
            else:
                pattern = r"\b" + re.escape(alias) + r"\b"

            if re.search(pattern, combined):
                found[alias] = canonical

        # Deduplicate by canonical name
        unique_skills = list(set(found.values()))
        unique_skills.sort()
        return unique_skills

    # ------------------------------------------------------------------
    # Experience Extraction
    # ------------------------------------------------------------------

    def _extract_experience(self, text: str, sections: Dict[str, str]) -> List[Dict[str, Any]]:
        """Extract work experience entries with company, role, duration, description."""
        exp_text = sections.get("experience", "")
        if not exp_text:
            return []

        entries = []
        lines = exp_text.strip().split("\n")
        current_entry: Optional[Dict[str, Any]] = None
        description_lines: List[str] = []

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            # Check if this line is an experience header (company / role line)
            # Heuristic: contains a date range or company-like patterns
            has_date = bool(YEAR_RANGE_PATTERN.search(stripped))
            is_bullet = stripped.startswith(("•", "-", "*", "–", "▪"))

            if has_date and not is_bullet:
                # Save previous entry
                if current_entry:
                    current_entry["description"] = " ".join(description_lines).strip()
                    entries.append(current_entry)
                    description_lines = []

                # Parse this header line
                date_match = YEAR_RANGE_PATTERN.search(stripped)
                duration = date_match.group(0) if date_match else ""
                header_text = YEAR_RANGE_PATTERN.sub("", stripped).strip(" |,–—-")

                # Try to split "Role at Company" or "Company — Role"
                role, company = self._split_role_company(header_text)

                current_entry = {
                    "company": company,
                    "role": role,
                    "duration": duration,
                    "description": "",
                }
            elif is_bullet and current_entry:
                description_lines.append(stripped.lstrip("•-*–▪ "))
            elif current_entry and not has_date:
                # Could be a continuation of the role/company or a description line
                if not current_entry.get("role") or not current_entry.get("company"):
                    # Maybe this is the second line of the header
                    role, company = self._split_role_company(stripped)
                    if role and not current_entry.get("role"):
                        current_entry["role"] = role
                    if company and not current_entry.get("company"):
                        current_entry["company"] = company
                else:
                    description_lines.append(stripped)

        # Don't forget the last entry
        if current_entry:
            current_entry["description"] = " ".join(description_lines).strip()
            entries.append(current_entry)

        return entries

    def _split_role_company(self, text: str) -> tuple:
        """
        Try to split a line like "Software Engineer at Google" or "Google — SDE Intern"
        into (role, company).
        """
        # Try splitting on common separators
        for sep in [" at ", " @ ", " — ", " – ", " - ", " | "]:
            if sep in text:
                parts = text.split(sep, 1)
                return parts[0].strip(), parts[1].strip()

        # If no separator, return full text as role, empty company
        return text.strip(), ""

    # ------------------------------------------------------------------
    # Projects Extraction
    # ------------------------------------------------------------------

    def _extract_projects(self, text: str, sections: Dict[str, str]) -> List[Dict[str, Any]]:
        """Extract project entries with title, description, technologies."""
        proj_text = sections.get("projects", "")
        if not proj_text:
            return []

        entries = []
        lines = proj_text.strip().split("\n")
        current_project: Optional[Dict[str, Any]] = None
        description_lines: List[str] = []

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            is_bullet = stripped.startswith(("•", "-", "*", "–", "▪"))

            # Heuristic for project title lines:
            # Not a bullet, not too long, often contains | or — for tech stack
            is_title = (
                not is_bullet
                and len(stripped) < 120
                and (
                    "|" in stripped
                    or "—" in stripped
                    or "–" in stripped
                    or bool(YEAR_RANGE_PATTERN.search(stripped))
                    or (len(stripped.split()) <= 10 and stripped[0].isupper())
                )
            )

            if is_title and not is_bullet:
                # Save previous project
                if current_project:
                    current_project["description"] = " ".join(description_lines).strip()
                    entries.append(current_project)
                    description_lines = []

                # Parse title line — split on | or — to separate title from tech stack
                parts = re.split(r"\s*[\|]\s*|\s*[—–]\s*", stripped, maxsplit=1)
                title = YEAR_RANGE_PATTERN.sub("", parts[0]).strip(" ,")
                tech_str = parts[1] if len(parts) > 1 else ""
                # Extract tech from the tech string
                technologies = [t.strip() for t in tech_str.split(",") if t.strip()] if tech_str else []

                current_project = {
                    "title": title,
                    "description": "",
                    "technologies": technologies,
                }
            elif is_bullet and current_project:
                description_lines.append(stripped.lstrip("•-*–▪ "))
            elif current_project:
                description_lines.append(stripped)

        # Don't forget the last project
        if current_project:
            current_project["description"] = " ".join(description_lines).strip()
            entries.append(current_project)

        return entries


# ---------------------------------------------------------------------------
# Convenience wrapper (backward-compatible with existing endpoint usage)
# ---------------------------------------------------------------------------
_parser_instance: Optional[ResumeParser] = None

def get_parser() -> ResumeParser:
    """Singleton accessor for ResumeParser."""
    global _parser_instance
    if _parser_instance is None:
        _parser_instance = ResumeParser()
    return _parser_instance
